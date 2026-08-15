from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

SPLUNK_GREEN = RGBColor(0x65, 0xA3, 0x00)
SPLUNK_DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)


def title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = MID_GRAY
    p.paragraph_format.space_after = Pt(18)
    return p


def heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_GREEN
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "65A300")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, size=10.5, bold=False, italic=False, space_after=8, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def checkbox(text, note=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run("\u2610  ")
    run.font.size = Pt(size + 1)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_GREEN
    run2 = p.add_run(text)
    run2.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.15)
    if note:
        p2 = doc.add_paragraph()
        run3 = p2.add_run(note)
        run3.font.size = Pt(9)
        run3.font.italic = True
        run3.font.color.rgb = MID_GRAY
        p2.paragraph_format.left_indent = Inches(0.45)
        p2.paragraph_format.space_after = Pt(8)
    return p


# --- Document content ---

title("Hyper-V Observability \u2014 Environment Verification Checklist")
subtitle(
    "Environment-specific checks to run in the customer's real fleet to tighten "
    "the solution before/during production rollout \u2014 2026-08-11"
)

body(
    "Everything below was validated on a live Azure test VM (nested Hyper-V) or "
    "confirmed empirically during the original POC, but has not yet been re-verified "
    "against the customer's actual production environment. None of these are blockers "
    "to the conversation \u2014 they're the concrete follow-ups to walk through together.",
    space_after=10,
)

# =====================================================================
heading1("Tier 0 \u2014 hyperv-scvmm-poller (never live-tested against a real SCVMM server)")

checkbox(
    "Run scvmm-poller against the real SCVMM console box and confirm hyperv_vm_up / "
    "hyperv_host_up populate correctly for every real power state (running, off, saved, paused).",
    "Gap #1. Code-complete but only unit-tested \u2014 this is the biggest untested surface area.",
)
checkbox(
    "Confirm the VirtualMachineManager PowerShell module is installed on their SCVMM "
    "console box, and which version.",
    "internal/scvmm shells out to it; never confirmed against their box.",
)
checkbox(
    "Verify guest_os classification accuracy against their real fleet's actual OS mix "
    "(SCVMM OperatingSystem field coverage, Secure Boot template fallback, naming heuristic "
    "false positives/negatives).",
    "Gap #8. Same code-complete-but-unvalidated caveat as gap #1.",
)

# =====================================================================
heading1("Gap #7 \u2014 network series (their own original finding)")

checkbox(
    "Re-check Get-VMNetworkAdapter / Get-Counter counter names on their real hosts to "
    "confirm the \u201cBytes/sec vs. Bytes Total/sec\u201d fix actually resolves it there.",
    "The fix was validated on a from-scratch Azure test host, not their cluster, which may "
    "be on a different collector/config version than the one that produced the original "
    "~20%-missing finding.",
)

# =====================================================================
heading1("Gap #9 \u2014 VMMS migration failures")

checkbox(
    "Verify the count connector's attributes[\"winlog.event_id\"] == 21026 condition "
    "against a raw ingested event for the exact collector version they deploy.",
    "The windowseventlogreceiver attribute key naming has changed across contrib releases.",
)

# =====================================================================
heading1("Gaps #3 / #6 \u2014 storage & naming (new findings from Azure nested-Hyper-V testing)")

checkbox(
    "Audit whether their VMs' disk files actually follow the New-VM default (VHDX filename "
    "matches owning VM name).",
    "If they use cloned templates, renamed disks, or reused VHDX files, storage metrics will "
    "silently misattribute to the wrong VM \u2014 a worse failure mode than the accepted "
    "19\u201323% unattributed residual, because it fails silently instead of visibly.",
)
checkbox(
    "Audit for duplicate VM names across their actual fleet.",
    "Both CPU/memory/storage (suffix-based) and network (GUID-based) counters collapse "
    "duplicate-named VMs onto one Splunk entity. This is a customer-side naming-hygiene fix, "
    "not something a collector processor can fix \u2014 worth surfacing now.",
)
checkbox(
    "Confirm the 77\u201381% VHD match rate holds on their production hosts.",
    "The live test pass only measured 9/18 (50%) on 3 synthetic VMs; the real ratio depends "
    "on how many DVD/ISO/pass-through disks they actually run.",
)

# =====================================================================
heading1("Tier 1.5 guest probe go/no-go (the big decision item)")

checkbox(
    "Run an Invoke-Command -VMId session latency/load test at their real VM density per host.",
    "Not yet tested beyond a 1\u20132 VM box.",
)
checkbox(
    "Audit Guest Integration Services version/coverage across their actual fleet.",
    "Older/legacy guests may lack it or run an outdated version.",
)
checkbox(
    "Complete a security review of the shared guest credential the probe uses "
    "(provisioning process, rotation policy).",
)

# =====================================================================
heading1("Credentials & deployment mechanics")

checkbox(
    "Confirm the account the Windows Service actually runs as in production (LocalSystem vs. "
    "a dedicated service account) matches the account under which cmdkey /generic:... was "
    "used to store the SCVMM/Splunk credentials.",
    "Windows Credential Manager entries are only readable by the exact account that set "
    "them \u2014 a mismatch here fails silently at runtime.",
)
checkbox(
    "Test the WiX MSI install/upgrade/uninstall path through whatever software-distribution "
    "tooling they actually use (GPO/SCCM), not just a manual install.",
)
checkbox(
    "Re-verify realm/ingest-token pairing in production before rollout.",
    "The test pass hit a real 401 from an org-token/realm mismatch \u2014 worth a deliberate "
    "check rather than discovering it live.",
)

# =====================================================================
heading1("Cutover")

checkbox(
    "Follow the shadow \u2192 diff \u2192 cutover plan per service on a small pilot set of hosts "
    "first (mixed VM density, at least one host with known pass-through/ISO disks) before "
    "fleet-wide rollout and before disabling the old scheduled tasks.",
)

doc.save("/Users/mbui/Desktop/Hyperv_O11y_Customer_Environment_Verification_Checklist.docx")
print("Saved to /Users/mbui/Desktop/Hyperv_O11y_Customer_Environment_Verification_Checklist.docx")
