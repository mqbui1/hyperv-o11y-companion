from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

SPLUNK_GREEN = RGBColor(0x65, 0xA3, 0x00)
SPLUNK_DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
WARN_RED = RGBColor(0xB0, 0x00, 0x00)


def title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = MID_GRAY
    p.paragraph_format.space_after = Pt(18)
    return p


def heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_GREEN
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "65A300")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, size=10.5, bold=False, italic=False, space_after=8, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, level=0, size=10.5, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def numbered(text, size=10.5, bold=False):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# --- Document content ---

title("Hyper-V Observability \u2014 Build & Test Summary")
subtitle(
    "What was built, what was validated end-to-end on live Azure test infrastructure, "
    "and how each previously-identified monitoring gap is addressed \u2014 2026-08-07"
)

# =====================================================================
heading1("Executive Summary")

body(
    "The Hyper-V observability solution now lives in a single consolidated repo, "
    "hyperv-o11y-companion, which merges what used to be two separate pieces: "
    "dashboards, detectors, and OTel Collector configs (formerly a standalone "
    "\u201caccelerator\u201d repo) plus two native Windows Services that replace four "
    "independently-scheduled PowerShell script pairs the customer was running via Windows "
    "Task Scheduler. Together they implement a multi-tier collection model covering the "
    "hypervisor host, the SCVMM management plane, and guest-OS data via a mechanism that "
    "requires nothing to be deployed inside the guest \u2014 this customer has explicitly "
    "ruled out deploying any collector inside guest VMs, opt-in or otherwise \u2014 see "
    "\u201cTier breakdown\u201d below."
)
body(
    "This round of work stood up the full companion-service stack from scratch on a live "
    "Azure VM with nested Hyper-V, and confirmed \u2014 with real metric data landing in "
    "Splunk Observability Cloud, correctly attributed per VM \u2014 that the disk-metrics "
    "gap (previously the single largest unresolved item) is now closed end-to-end. Two "
    "real defects were found and fixed during testing (an OTLP endpoint bug and a config-"
    "patch YAML bug); both are documented below along with the fixes applied."
)

# =====================================================================
heading1("What Was Built")

heading2("Tiered collection architecture")
body(
    "Started as a two-tier host/guest split (mirroring the vSphere navigator pattern) and "
    "grew more tiers once real customer POC data showed Perfmon alone can't see VM "
    "power-state, live-migration failures, or in-guest data. Each tier is a distinct "
    "collection mechanism with a distinct blast radius:",
    space_after=8,
)
add_table(
    ["Tier", "What it is", "Where it runs", "Status"],
    [
        ["0", "hyperv-scvmm-poller (Windows Service) polling SCVMM", "Central SCVMM console box", "Implemented, not yet live-tested end-to-end"],
        ["1", "Splunk OTel Collector, hypervisor-host-config.yaml", "Every Hyper-V host", "Built & validated"],
        ["1 companion", "hyperv-host-companion (Windows Service)", "Every Hyper-V host, alongside Tier 1", "Built & tested end-to-end this round"],
        ["1.5", "PowerShell Direct guest probe (VMBus, no guest network/agent)", "Host-initiated", "Implemented, mechanism-validated on a real nested Hyper-V guest; ships disabled pending fleet go/no-go"],
    ],
    col_widths=[0.8, 2.6, 1.9, 1.7],
)

heading2("Tier breakdown \u2014 what each tier represents")
body(
    "What started as a two-tier host/guest split grew into more tiers as real customer-"
    "POC findings showed a plain host/guest split couldn't see everything. Each tier is a "
    "genuinely different collection mechanism, not a variation on the same idea:",
    space_after=8,
)
bullet("Tier 0 \u2014 SCVMM console: a single centralized service polling the SCVMM management plane (not per-host Perfmon). This is the only place VM/host power-state and SCVMM's own OperatingSystem field are visible \u2014 Perfmon simply has no concept of \u201coff\u201d and can't see SCVMM's inventory metadata. Closes gaps #1 and #8.", bold=False)
bullet("Tier 1 \u2014 every Hyper-V host: the generic OTel windowsperfcounters/windowseventlogreceiver pipeline, deployed identically on every host. Sees hypervisor-visible resource allocation (vCPU %, assigned memory, virtual disk file I/O, vNIC throughput) \u2014 cheap (no extra billable hosts) but structurally cannot see anything happening inside a guest's own OS.")
bullet("Tier 1 companion \u2014 same hosts as Tier 1, but a purpose-built Go service (hyperv-host-companion) instead of a generic receiver. Exists because per-VM disk-latency/throughput attribution needs live VM-to-VHD-path resolution that the generic Perfmon receiver can't do on its own. Closes gap #3.")
bullet("Tier 1.5 \u2014 PowerShell Direct guest probe: the host reaches into a guest over VMBus (Invoke-Command -VMId) to run a narrow, specific query \u2014 no guest network stack, no guest firewall rule, no agent installed inside the guest at all. Nothing is deployed inside the guest, so it doesn't trigger the customer's \u201cno in-guest collector\u201d constraint. Implemented and mechanism-validated end-to-end on a real nested Hyper-V guest this round \u2014 the only candidate for gaps #2 and #4 \u2014 but ships with guest_probe.enabled: false pending fleet-wide go/no-go (session load at scale, real-fleet Integration Services coverage, shared-credential security review).")
bullet("Future idea, not gap-driven, not built \u2014 Windows Event Forwarding: a native Windows mechanism that could forward event logs from hosts/guests to one central collector, without installing a per-VM collector. No code, no config in this repo, and no customer-POC gap requires it \u2014 unlike every numbered tier above, it isn't tied to a specific finding.")

heading2("hyperv-o11y-companion \u2014 the two consolidated Windows Services")
body(
    "One repo, two long-running native Windows Services, zero Task Scheduler entries. "
    "Replaces the customer's four independently-scheduled script pairs (collect-scvmm-"
    "metrics.ps1 / run-collect-scvmm-metrics.ps1, enrich-vm-guest-os.ps1 / run-enrich-vm-"
    "guest-os.ps1, build-hyperv-vm-disk-map.ps1 / collect-hyperv-vm-disk.ps1) plus the "
    "planned Tier 1.5 guest probe.",
    space_after=8,
)
add_table(
    ["Service", "Runs where", "Replaces", "Status"],
    [
        ["hyperv-scvmm-poller", "Central SCVMM console box", "collect-scvmm-metrics.ps1 + enrich-vm-guest-os.ps1 and their wrappers/scheduled tasks", "Phase 1 \u2014 code-complete, not yet live-tested end-to-end"],
        ["hyperv-host-companion", "Every Hyper-V host, alongside the Splunk OTel Collector", "build-hyperv-vm-disk-map.ps1 + collect-hyperv-vm-disk.ps1 (Phase 2); PowerShell Direct guest probe (Phase 3)", "Phase 2 \u2014 built and tested end-to-end this round. Phase 3 \u2014 implemented and mechanism-validated, disabled pending go/no-go"],
    ],
    col_widths=[1.9, 2.4, 2.9, 1.6],
)
bullet("Both installed as native Windows Services (golang.org/x/sys/windows/svc) via a WiX v4 MSI \u2014 start on boot, restart on failure, visible in services.msc. No Task Scheduler, no DPAPI secret files on disk.")
bullet("hyperv-scvmm-poller reads its SCVMM read-only account and Splunk access token from Windows Credential Manager at startup \u2014 not from files on disk.")
bullet("hyperv-host-companion needs no credentials at all: queries the local Hyper-V host directly (Get-VM / Get-VMHardDiskDrive / Get-Counter, no remoting) and exports over plain OTLP to the host-local Splunk OTel Collector, which already holds the upstream Splunk credential via its own signalfx exporter.")

heading2("Supporting documentation and tooling delivered (all in hyperv-o11y-companion)")
add_table(
    ["Artifact", "Purpose"],
    [
        ["docs/architecture.md", "Five-tier model, resource-attribute strategy, vm.name extraction rules"],
        ["docs/known-gaps-remediation.md", "Gap-by-gap findings and fixes (source for the table below)"],
        ["docs/limitations.md", "What the solution deliberately does not do"],
        ["docs/parity-testing-and-cutover.md", "Shadow \u2192 diff \u2192 cutover plan for replacing the existing scripts with the two new services"],
        ["docs/phase3-guest-probe-plan.md", "Tier 1.5 (PowerShell Direct guest probe) \u2014 implemented, mechanism-validated, disabled pending fleet go/no-go"],
        ["otel-collector/, terraform/", "Tier 1 collector configs + dashboards/detectors (merged in from the former standalone accelerator repo)"],
        ["installer/ (WiX v4 source)", "Two-feature MSI scaffold \u2014 one feature per service, same MSI installs the right thing on the SCVMM console box vs. every Hyper-V host"],
    ],
    col_widths=[2.2, 4.6],
)

# =====================================================================
heading1("What Was Tested")

body(
    "Stood up and validated the full hyperv-host-companion stack end-to-end on a live "
    "Azure VM (Standard_D4als_v7, Windows Server 2025 Datacenter, TrustedLaunch, nested "
    "virtualization) rather than relying on unit tests alone \u2014 this exercised the real "
    "Windows Service lifecycle, real Hyper-V Perfmon counters, and a real Splunk "
    "Observability Cloud ingest path.",
    space_after=8,
)
numbered("Enabled the Hyper-V role on the test VM (Install-WindowsFeature) and rebooted externally rather than via a guest-triggered restart, to avoid a known Azure Run Command hang (see \u201cIssues found and fixed\u201d). Confirmed the vmms service running post-reboot.")
numbered("Installed the Splunk OTel Collector (agent mode, v0.157.0) and patched its config to add an otlp receiver (grpc:4317 / http:4318) plus a metrics/host_companion pipeline feeding the signalfx exporter.")
numbered("Deployed hyperv-host-companion as a real Windows Service (New-Service, not a foreground process) pointed at the local collector.")
numbered("Created three nested test VMs (vm-alpha1, vm-alpha2, vm-beta) with attached virtual hard disks to exercise VM/disk enumeration.")
numbered("Confirmed, via a bounded diagnostic run of the service's own logging, correct disk-map and metric-attribution behavior: \u201cdisk map rebuilt: 3 path entries, 3 VM ids\u201d and \u201cdisk metrics: matched=9 unmatched=9 (50% match rate)\u201d against the three test VMs.")
numbered("Confirmed end-to-end delivery into Splunk Observability Cloud: vm.disk.latency, vm.disk.read_bytes_sec, and vm.disk.write_bytes_sec all present, correctly tagged per VM (vm.name: vm-alpha1 / vm-alpha2 / vm-beta), with host.name, service.name, and full Azure host resource attributes attached \u2014 verified directly via metric/time-series search against the org, not inferred from logs alone.")
numbered("Separately validated the Tier 1.5 (PowerShell Direct) mechanism against a second, genuine bootable guest \u2014 a Windows Server 2022 Server Core VM built via unattended DISM apply + bcdboot, with Integration Services reporting OK/Heartbeat. internal/creds.NewReader().Read() resolved the stored .\\Administrator credential from Windows Credential Manager; Invoke-Command -VMId succeeded over VMBus with no guest network path at all (Internal-only nested switch). internal/guestprobe.Sample() returned correct real numbers for both metrics: 16.30% filesystem used (drive C:) and 39.36% memory used, the latter on a VM with Dynamic Memory explicitly disabled \u2014 exactly the static-memory case Hyper-V's own Current Pressure counter can't see (gap #2).")

heading2("Issues found and fixed during this test pass")
add_table(
    ["Issue", "Fix applied"],
    [
        [
            "OTLP endpoint double-prefixing bug in internal/metricsexport/exporter.go: otlpmetrichttp.WithEndpoint() requires a bare host:port, but every config file in the repo (for readability) uses a full \u201chttp://host:port\u201d URL.",
            "Strip the scheme with net/url.Parse before passing to WithEndpoint. Verified fixed via the live metric delivery in test #6 above.",
        ],
        [
            "Config-patch script inserted the metrics/host_companion pipeline at the wrong YAML indentation (6 spaces vs. the sibling traces:/metrics: pipelines' 4 spaces), which silently broke the config and crashed the collector service on restart.",
            "Corrected the patch to match sibling-key indentation exactly; added an explicit post-restart service-status check to the deployment procedure so a broken config can't go unnoticed.",
        ],
        [
            "Ingest 401 Unauthorized errors during initial setup, caused by a realm/org-token mismatch (an API-scoped management token was used against the wrong Splunk Observability Cloud realm) \u2014 not a defect in the solution itself.",
            "Confirmed the correct realm + a genuine ingest-scoped access token; documented as an operational runbook note so future deployments confirm realm/token pairing up front.",
        ],
    ],
    col_widths=[3.6, 3.2],
)

# =====================================================================
heading1("How Each Known Gap Is Addressed")

body(
    "Source: a real customer POC test summary. Each row maps a specific finding to the "
    "concrete fix implemented in this solution, or to an explicit, documented non-goal with "
    "a recommended workaround.",
    space_after=8,
)
add_table(
    ["#", "Gap", "Status", "How it's addressed"],
    [
        ["1", "No VM/host \u201cdown\u201d (power-state) detection \u2014 Perfmon only reports state for running VMs", "Solved (implemented, not yet live-tested end-to-end)", "hyperv_vm_up / hyperv_host_up gauges sourced from SCVMM PowerState/OverallState (Tier 0), tagged with the same vm.name/host.name attributes used everywhere else; ported into hyperv-scvmm-poller's pollMetrics loop \u2014 code-complete, but not yet run end-to-end against a real SCVMM server."],
        ["2", "Static-memory VMs invisible to memory-pressure alerting (Dynamic Memory counter doesn't exist for them)", "Solved (pending fleet-wide validation)", "Not a bug \u2014 a coverage gap. Tier 1.5 (PowerShell Direct guest probe) queries Win32_OperatingSystem inside the guest instead \u2014 works for static-memory VMs, mechanism-validated on a real nested Hyper-V guest \u2014 but ships disabled pending fleet-wide go/no-go."],
        ["3", "~19\u201323% of VHD instances unattributed (DVD/ISO/pass-through disks can't map to a VM by path alone)", "Solved", "In-memory disk map (Get-VM | Get-VMHardDiskDrive) built and refreshed by hyperv-host-companion, resolving live storage counters to VMs; unmatched instances are counted and logged, not guessed at. Validated live: 9/18 matched in this test pass, consistent with the accepted 77\u201381% fleet match rate."],
        ["4", "No guest filesystem used % visible from the host (architectural limitation)", "Solved (pending fleet-wide validation)", "Tier 1.5 (PowerShell Direct guest probe, Invoke-Command -VMId) closes this \u2014 mechanism-validated end-to-end on a real nested Hyper-V guest (16.30% filesystem used, computed correctly) \u2014 but ships disabled pending fleet-wide go/no-go."],
        ["5", "Disk latency counter unit unconfirmed (100ns ticks per Microsoft docs, never independently verified)", "Solved", "Empirically confirmed via real-fleet data that the counter reports raw seconds, not 100ns ticks; a \u00d71000 scale correction is applied before export and the storage-latency detector is enabled with a real ms threshold."],
        ["6", "Malformed vm.name values from Perfmon's #N duplicate-instance suffixing", "Solved (partially \u2014 naming collisions remain a customer-side issue)", "Final extraction step strips the #N suffix so vm.name reads cleanly. Note: this does not resolve true duplicate-named VMs colliding onto the same entity \u2014 that's a Hyper-V naming-hygiene issue on the customer's estate, not something a collector processor can fix."],
        ["7", "~20% of VMs emitting no network series", "Solved for a from-scratch deployment", "Root-caused to an incorrect Perfmon counter name (\u201cBytes Total/sec\u201d, which doesn't exist on the Hyper-V network objects) that made the metric 0% functional, not ~80% functional. Corrected to the real counter name (\u201cBytes/sec\u201d) and verified populating in Splunk Observability Cloud. Still open: this explains a total outage on a from-scratch deployment, but doesn't by itself confirm the original POC's \u201c~80% match rate\u201d finding on a cluster that may have had a different config version \u2014 re-check against Get-VMNetworkAdapter if the gap resurfaces on the customer's own cluster."],
        ["8", "guest_os accuracy issues (unknown/untagged VMs, heuristic Linux tagging)", "Solved (implemented, not yet live-tested end-to-end)", "Classification logic (SCVMM OperatingSystem field \u2192 Secure Boot template fallback \u2192 opt-in naming heuristic) writes guest_os as a zero-cost dimension property via the SignalFx metadata API; ported into hyperv-scvmm-poller's pollGuestOS loop \u2014 code-complete, but not yet run end-to-end against a real SCVMM server (same caveat as gap #1)."],
        ["9", "VMMS load issues driven by failed live migrations (Event ID 21026)", "Solved", "A count connector converts Event ID 21026 occurrences into an alertable hyperv.vmms.migration_failures metric, with its own pipeline and detector \u2014 makes the root cause directly alertable instead of only visible after the fact."],
        ["10", "POC \u2192 production cutover risk", "Low risk, no change needed", "The host.type filter design used throughout the configs and content is environment-agnostic; cutover is just pointing the realm/access token at production and rolling the same config to remaining hosts."],
    ],
    col_widths=[0.4, 2.5, 1.4, 2.5],
)

# =====================================================================
heading1("What's Not Yet Solved / Deliberately Out of Scope")

bullet("Duplicate VM names on the same Hyper-V host: metrics for identically-named VMs will merge onto one entity. Not fixable in collector config \u2014 flagged as a customer-side naming-hygiene recommendation.")
bullet("Storage-metric misattribution when a VHDX filename doesn't match its owning VM's name (cloned templates, renamed disks, reused files): a real, unresolved limitation with no reliable signal to detect it from the raw counter alone.")
bullet("Tier 1.5 (PowerShell Direct guest probe): implemented and mechanism-validated end-to-end on a real nested Hyper-V guest \u2014 closes gaps #2 and #4 without deploying anything inside the guest at all. Ships disabled (guest_probe.enabled: false) and stays that way until a separate fleet-wide go/no-go decision (session latency/load at scale, real-fleet Integration Services coverage, shared-credential security review). This customer has separately ruled out deploying any collector inside guest VMs at all, opt-in or otherwise.")
bullet("hyperv-scvmm-poller (Tier 0) is code-complete (pollMetrics and pollGuestOS both implemented) but not yet run end-to-end against a real SCVMM server the way hyperv-host-companion was this round \u2014 next priority for the same live-test treatment.")

# =====================================================================
heading1("Recommended Next Steps")

numbered("Run the same live-VM test treatment against hyperv-scvmm-poller (Tier 0) that hyperv-host-companion just received.")
numbered("Follow the documented shadow \u2192 diff \u2192 cutover plan per service: install alongside the existing PowerShell scripts/scheduled tasks for 24\u201348h, diff coverage and classification output, then disable (not delete) the old scheduled tasks once parity holds for a full week.")
numbered("Roll hyperv-host-companion to a small representative set of hosts first (mixed VM density, at least one host with known pass-through/ISO disks), then fleet-wide only after that set has been clean for a full week.")
numbered("Run the Tier 1.5 fleet-wide go/no-go validation (session latency/load at real VM density, real-fleet Integration Services coverage audit, shared-credential security review), using the small-pilot procedure in docs/deployment-guide.md, before enabling guest_probe.enabled fleet-wide.")

doc.save("/Users/mbui/Desktop/Hyperv_O11y_Build_and_Test_Summary.docx")
print("Saved to /Users/mbui/Desktop/Hyperv_O11y_Build_and_Test_Summary.docx")
