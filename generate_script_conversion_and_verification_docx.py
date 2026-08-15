from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

SPLUNK_GREEN = RGBColor(0x65, 0xA3, 0x00)
SPLUNK_DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
WARN_RED = RGBColor(0xB0, 0x00, 0x00)


def title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = MID_GRAY
    p.paragraph_format.space_after = Pt(18)
    return p


def heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_GREEN
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "65A300")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = SPLUNK_DARK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, size=10.5, bold=False, italic=False, space_after=8, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, level=0, size=10.5, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def numbered(text, size=10.5, bold=False):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# =====================================================================
# Document content
# =====================================================================

title("Hyper-V Observability \u2014 Script Conversion & Guest-Probe Mechanics")
subtitle(
    "How each original PowerShell script maps onto the hyperv-o11y-companion Go services, "
    "and exactly how Tier 1.5 fetches data from inside a guest VM \u2014 2026-08-07"
)

# =====================================================================
heading1("Part 1 \u2014 PowerShell Script \u2192 Go Service Conversion")

body(
    "The customer originally ran four independently-scheduled PowerShell script pairs via "
    "Windows Task Scheduler \u2014 each pair a \u201cdo the work\u201d script plus a thin wrapper the "
    "scheduled task actually invoked. hyperv-o11y-companion replaces all four pairs with two "
    "long-running native Windows Services, each running the equivalent logic as goroutine "
    "tickers instead of scheduled one-shot invocations. No Task Scheduler entries, no DPAPI "
    "secret files on disk in either case."
)

add_table(
    ["Original script pair", "Go service", "Go package(s) / function", "Scheduling equivalent"],
    [
        ["collect-scvmm-metrics.ps1\n+ run-collect-scvmm-metrics.ps1",
         "hyperv-scvmm-poller",
         "internal/scvmm (Client.Hosts/VMs)\ncmd/scvmm-poller/main.go: pollMetrics()",
         "metricsTicker (cfg.metrics.interval, default 60s) inside runService's for/select loop"],
        ["enrich-vm-guest-os.ps1\n+ run-enrich-vm-guest-os.ps1",
         "hyperv-scvmm-poller",
         "internal/guestos (Classify/NameHeuristic)\ninternal/metadata (SetGuestOS)\ncmd/scvmm-poller/main.go: pollGuestOS()",
         "guestOSTicker (cfg.guest_os.interval, default 1h) \u2014 independent of the metrics ticker, same service"],
        ["build-hyperv-vm-disk-map.ps1",
         "hyperv-host-companion",
         "internal/hyperv.BuildDiskMap\ncmd/host-companion/main.go: buildDiskMap()",
         "buildTicker (cfg.disk_map.build_interval, default 1h)"],
        ["collect-hyperv-vm-disk.ps1",
         "hyperv-host-companion",
         "internal/hyperv.SampleStorageCounters\ninternal/diskattr.Resolve\ncmd/host-companion/main.go: sampleAndExport()",
         "sampleTicker (cfg.disk_metrics.sample_interval, default 60s)"],
    ],
    col_widths=[1.7, 1.5, 2.3, 1.8],
)

heading2("1. collect-scvmm-metrics.ps1 \u2192 hyperv-scvmm-poller (pollMetrics)")
bullet("Original: scheduled task connects to SCVMM, runs Get-SCVMHost / Get-SCVirtualMachine, computes host/VM up-down state and host memory used, emits metrics.")
bullet("Go port: internal/scvmm.Client.Hosts()/VMs() shell out to a single powershell.exe process per poll cycle running the identical VMM cmdlets (Get-SCVMHost, Get-SCVirtualMachine) with the same field selection and cluster Where-Object filter \u2014 SCVMM has no on-prem REST API and no Go SDK for the VirtualMachineManager module, so this remains the one deliberate PowerShell touchpoint in the service.")
bullet("pollMetrics() (cmd/scvmm-poller/main.go) applies the identical up/down logic (Overall == \"Responding\"/\"OK\" \u2192 host up; Status == \"Running\" \u2192 VM up) and records hyperv_host_up, hyperv_vm_up, and hyperv_host_memory_used_mb as OTel Float64Gauge instruments, tagged with host.name / vm.name / hypervisor.host.name \u2014 exported via OTLP, not written to a file for another script to pick up.")
bullet("Credential handling changed: the SCVMM account's username/password are read once at service startup from Windows Credential Manager (internal/creds) and passed into each PowerShell invocation via process environment variables (HYPERV_O11Y_SCVMM_*), built into a PSCredential from SecureString entirely in-process \u2014 never written to a DPAPI file on disk and never present in argv/process listing.")
bullet("Scheduling changed: what was a Task-Scheduler-invoked script + wrapper is now one ticker (time.NewTicker) inside a single always-running Windows Service, with an immediate run on startup before the first interval tick.")

heading2("2. enrich-vm-guest-os.ps1 \u2192 hyperv-scvmm-poller (pollGuestOS)")
bullet("Original: separate scheduled task with its own SCVMM connection; classifies guest_os from the OperatingSystem field (Get-GuestOs), falls back to a Secure Boot template check for still-unknown Gen2 VMs, then an opt-in naming heuristic (Get-NameHeuristicOs) as a last resort; writes the result as a guest_os dimension property via the SignalFx metadata API (GET \u2192 merge \u2192 PUT); skips (does not clobber) ambiguous duplicate VM names.")
bullet("Go port: pollGuestOS() reuses the same client.VMs() call already made by pollMetrics \u2014 one shared SCVMM client, two independent tickers in one service, instead of two scripts each opening their own SCVMM connection. internal/guestos.Classify() ports the exact regex classification rules (unknown/windows/linux/other); internal/guestos.NameHeuristic() ports the substring-match fallback, gated by cfg.guest_os.name_heuristic, returning \"\" (no guess) rather than a false positive when nothing matches \u2014 same conservative behavior as the original. internal/metadata.Client.SetGuestOS() ports the GET\u2192merge\u2192PUT dimension-API pattern against the vm dimension, including the same changed/unchanged bookkeeping the original script logs.")
bullet("Duplicate-name handling preserved exactly: pollGuestOS() counts VM-name occurrences first and skips (skipped++, continue) any name seen more than once, matching the original script's \u201cskip, don't clobber\u201d behavior for ambiguous duplicates.")
body(
    "Verification note: config/config.go defines a guest_os.secure_boot_fallback flag "
    "(GuestOSConfig.SecureBootFallback), but pollGuestOS() in the current cmd/scvmm-poller/"
    "main.go only calls guestos.Classify() and, if still unknown, guestos.NameHeuristic() "
    "\u2014 the Secure Boot template fallback step described in known-gaps-remediation.md's "
    "original-script writeup does not yet have a corresponding function wired into the poll "
    "loop. The config flag is scaffolded; the fallback logic itself is not yet implemented "
    "in internal/guestos as read for this document.",
    italic=True, color=WARN_RED, space_after=10,
)

heading2("3. build-hyperv-vm-disk-map.ps1 \u2192 hyperv-host-companion (buildDiskMap)")
bullet("Original: periodic scheduled task runs Get-VM | Get-VMHardDiskDrive and writes a VHD-path \u2192 VM-name JSON cache file for a second script to read, with a configurable timeout (-TimeoutSec) since VMMS can be slow under load.")
bullet("Go port: internal/hyperv.BuildDiskMap(ctx) runs the identical Get-VM | ForEach-Object { Get-VMHardDiskDrive ... } query locally, decodes the JSON directly into a DiskMap struct (ByPath and ByID indexes, path-normalized to lowercase/cleaned), and returns it in-process \u2014 no intermediate JSON file. buildDiskMap() in cmd/host-companion/main.go wraps this with a context timeout from cfg.disk_map.build_timeout, and on failure keeps the previous map in place rather than blocking or exporting unattributed metrics, replicating the original script's timeout-fallback behavior.")
bullet("Handoff mechanism changed: the original two-script JSON-cache-file handoff is replaced by an in-memory diskMapState struct (a sync.RWMutex-guarded pointer) shared between the builder ticker and the sampler ticker inside the same process \u2014 no file I/O, no risk of reading a half-written cache file.")

heading2("4. collect-hyperv-vm-disk.ps1 \u2192 hyperv-host-companion (sampleAndExport)")
bullet("Original: separate scheduled task samples the Hyper-V Virtual Storage Device Perfmon counters directly via Get-Counter (not the OTel windowsperfcounters receiver, so it can resolve VM attribution before export), reads the JSON disk-map cache file written by the builder script, resolves each instance to a VM (path-suffix match with a .vmgs VM-Id fallback), and applies an empirically-confirmed seconds\u2192ms latency scale correction before emitting metrics.")
bullet("Go port: internal/hyperv.SampleStorageCounters(ctx) runs the identical live Get-Counter collection against the same three counter paths (Latency, Read Bytes/sec, Write Bytes/sec) in one PowerShell invocation. internal/diskattr.Resolve() ports the resolution logic exactly: exact path match first, then a suffix match (handles UNC vs. local path prefix differences to the same file), then a .vmgs \u2192 VM-GUID fallback against the ID index. sampleAndExport() in cmd/host-companion/main.go applies the same \u00d71000 (seconds\u2192ms) latency scale correction (cfg.disk_metrics.latency_scale, default 1000.0 \u2014 see gap #5) and records vm.disk.latency / vm.disk.read_bytes_sec / vm.disk.write_bytes_sec as OTel gauges, exported via OTLP to the host-local Splunk OTel Collector.")
bullet("Match-rate accounting preserved: unmatched instances (DVD/ISO/pass-through disks with no reliable VM signal in the raw path) are counted and logged (\u201cdisk metrics: matched=%d unmatched=%d\u201d), not guessed at \u2014 same accepted ~19\u201323% residual gap (77\u201381% fleet match rate) as the original script pair, per known-gaps-remediation.md gap #3.")

heading2("Cross-cutting conversions (not specific to one script)")
add_table(
    ["Concern", "Original approach", "Go service approach"],
    [
        ["Credential storage", "DPAPI-encrypted files under C:\\ProgramData\\O11yScripts\\*.pw.txt", "Windows Credential Manager generic credentials (internal/creds), one secret per credential, set once via cmdkey, readable only by the service account, resolved at process startup \u2014 never written to disk in this repo's own code"],
        ["Scheduling", "Windows Task Scheduler entries invoking one-shot script + wrapper pairs", "time.Ticker loops inside two always-running Windows Services (golang.org/x/sys/windows/svc), installed via a WiX v4 MSI \u2014 start on boot, restart on failure, visible in services.msc"],
        ["Inter-process handoff", "JSON cache files on disk (disk map) written by one script, read by another", "In-memory shared state within one process (sync.RWMutex-guarded struct, or a single shared client call reused by two tickers) \u2014 no file I/O, no partial-write race"],
        ["SCVMM/PowerShell access", "Each script opens its own SCVMM connection and/or local PowerShell session", "One narrow shell-out adapter per package (internal/scvmm, internal/hyperv, internal/guestprobe) \u2014 same underlying cmdlets, called from a long-running process instead of a per-run script invocation"],
    ],
    col_widths=[1.6, 2.7, 3.0],
)

# =====================================================================
heading1("Part 2 \u2014 How Guest VM Data Is Fetched (Tier 1.5)")

body(
    "Tier 1.5 (internal/guestprobe, wired into hyperv-host-companion) is the only mechanism "
    "in this solution that reads data from inside a guest VM's own OS \u2014 without deploying "
    "any collector or agent inside the guest. It closes gap #2 (static-memory VM memory "
    "pressure) and gap #4 (guest filesystem used %). The mechanism is PowerShell Direct "
    "(Invoke-Command -VMId), which tunnels a PowerShell session from the host into a guest "
    "over VMBus \u2014 Hyper-V's host\u2194guest integration-component channel \u2014 with no guest "
    "network stack, no guest-facing firewall rule, and no in-guest agent process required.",
    space_after=10,
)

heading2("End-to-end process, one sample cycle")
numbered("Trigger: a dedicated guestProbeTicker inside hyperv-host-companion's runService loop fires every cfg.guest_probe.sample_interval (default 5m). This ticker is only created at all if guest_probe.enabled is true AND vm_include is non-empty \u2014 otherwise the select statement's channel case is a permanently-nil channel, cleanly disabling the whole tier with zero runtime overhead.")
numbered("VM enumeration: sampleGuestProbe() does NOT issue its own Get-VM call. It reuses the VM ID \u2192 VM name map already built by the disk-map builder ticker (internal/hyperv.BuildDiskMap, the same map Tier 1 companion's disk metrics use), filtering each entry's VMName against the vm_include glob patterns (filepath.Match, e.g. \"WebServer*\") \u2014 an empty include list means no VM is probed even when enabled.")
numbered("Credential retrieval: once per sample cycle, creds.NewReader().Read(cfg.guest_probe.credential_name) reads a single shared guest-local account from Windows Credential Manager (default name hyperv-o11y/guest-probe, provisioned via cmdkey /generic:...). This is a credential valid INSIDE the guest OS, not a host-level or domain-admin account \u2014 and it is one shared account for the fleet, not per-VM, which is exactly go/no-go criterion #3 in docs/phase3-guest-probe-plan.md (not yet resolved).")
numbered("Per-VM timeout: for each matching VM, a context.WithTimeout bounded by cfg.guest_probe.sample_timeout (default 30s) is created before calling guestprobe.Sample() \u2014 this prevents one guest with stale/missing Integration Services from hanging the entire probe cycle for every other VM on the host.")
numbered("The PowerShell Direct call: guestprobe.Sample(ctx, vmID, cred) shells to a single powershell.exe invocation. The guest account's username and VM ID are passed as command-line arguments; the password is passed via a process environment variable (HYPERV_O11Y_GUESTPROBE_PASS), never as a command-line argument, so it never appears in a process listing (ps / Get-Process command line).")
numbered("Inside that PowerShell process: a SecureString is built from the env-var password, wrapped into a PSCredential, and passed to Invoke-Command -VMId $VMId -Credential $cred -ScriptBlock { ... }. This is the actual VMBus tunnel into the guest \u2014 no TCP/IP path to the guest is used at all, which is why this works even on guests with no network connectivity configured.")
numbered("Inside the guest (the ScriptBlock, executing on the GUEST's own PowerShell, not the host's): Get-Volume | Where DriveType -eq Fixed | Select DriveLetter,Size,SizeRemaining gathers per-drive-letter filesystem capacity (gap #4), and Get-CimInstance Win32_OperatingSystem | Select TotalVisibleMemorySize,FreePhysicalMemory gathers the guest's own OS-reported memory usage (gap #2) \u2014 both facts in ONE Invoke-Command session, not two, since the session itself is the expensive/risky part at fleet scale (go/no-go criterion #1), not what's queried once inside it.")
numbered("Both results are packaged into one PSCustomObject and serialized with ConvertTo-Json -Depth 4 as the script's stdout. A finally block explicitly clears the password/SecureString/PSCredential variables (Remove-Variable) before the session ends.")
numbered("Back in Go: the JSON is decoded into a GuestSample{Filesystem []FilesystemSample, Memory MemorySample} struct. FilesystemSample.UsedPercent() computes (Size-SizeRemaining)/Size*100 per drive (skipping any volume with Size<=0 rather than dividing by zero); MemorySample.UsedPercent() computes the equivalent from TotalVisibleMemoryKB/FreePhysicalMemoryKB.")
numbered("Export: sampleGuestProbe() records vm.guest.filesystem.used_percent once per fixed drive letter (tagged vm.name + drive_letter) and vm.guest.memory.used_percent once per VM (tagged vm.name) as OTel Float64Gauge instruments, exported over the same OTLP path hyperv-host-companion already uses for its Tier 1 companion disk metrics \u2014 landing on the same Splunk Observability Cloud entity as that VM's Tier 1 metrics with no extra correlation work.")
numbered("Failure isolation: if any step for one VM fails (PowerShell error, timeout, missing Integration Services), it is logged (\u201cguest probe: vm=%s: %v\u201d) and that VM is skipped \u2014 the loop continues to the next VM in vm_include. One bad guest never blocks metrics for the rest of the fleet.")

heading2("Why this closes gap #2 specifically for static-memory VMs")
body(
    "Hyper-V's own \u201cCurrent Pressure\u201d counter (Hyper-V Dynamic Memory VM object, used by "
    "Tier 1) only exists for VMs with Dynamic Memory enabled \u2014 it is a Dynamic-Memory "
    "balloon-driver artifact, not a general memory-usage signal. vm.guest.memory."
    "used_percent instead asks the guest's own OS how much of ITS memory it's using "
    "(Win32_OperatingSystem), which works identically whether Dynamic Memory is on or off. "
    "This was validated on a real nested-Hyper-V guest with Dynamic Memory explicitly "
    "disabled \u2014 exactly the case Tier 1 structurally cannot see \u2014 and returned a correct "
    "39.36% memory-used figure.",
    space_after=10,
)

heading2("Live-migration behavior (no explicit cleanup logic needed)")
body(
    "Metrics are keyed by the vm.name dimension only, not a host+VM composite key. Because "
    "sampleGuestProbe() only probes VMs present in the LOCAL disk-map enumeration, and a "
    "Hyper-V VM is exclusively owned by one host at a time, a live migration is handled "
    "implicitly: the source host's next disk-map rebuild drops the VM from its own "
    "enumeration (its guest-probe ticker simply stops querying an ID that's now invalid "
    "there), and the destination host's next rebuild picks it up. The only gap is the window "
    "between migration completing and each host's next disk_map.build_interval tick "
    "(default 1h) \u2014 not yet tuned against a real migration-frequency profile.",
    space_after=6,
)

heading2("Requirements and constraints")
bullet("Guest Integration Services must be running and current inside the guest \u2014 confirmed per-VM via Get-VM <name> | Select IntegrationServicesState (must report \u201cUp to date\u201d) before including a VM in vm_include.")
bullet("Windows guests only \u2014 PowerShell Direct is a Windows-guest-only mechanism; Linux VMs needing this data are not covered by this mechanism.")
bullet("A single shared guest-local credential is assumed fleet-wide by the current implementation \u2014 whether that's acceptable to the customer's security review, or per-VM/per-domain credential handling is required instead, is go/no-go criterion #3 and is not yet resolved.")
bullet("Ships with guest_probe.enabled: false by default and an empty vm_include \u2014 no VM is probed until both are explicitly opted in, and should stay off fleet-wide until all three go/no-go criteria in docs/phase3-guest-probe-plan.md are validated.")

doc.save("/Users/mbui/Desktop/Hyperv_O11y_Script_Conversion_and_Guest_Probe_Guide.docx")
print("Saved to /Users/mbui/Desktop/Hyperv_O11y_Script_Conversion_and_Guest_Probe_Guide.docx")
